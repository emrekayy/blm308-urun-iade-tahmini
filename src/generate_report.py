"""BLM308 final raporunu DOCX ve PDF formatında oluşturur."""

from __future__ import annotations

import json
import subprocess
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import Image, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

PROJECT_ROOT = Path(__file__).resolve().parent.parent
if str(PROJECT_ROOT) not in sys.path:
    sys.path.insert(0, str(PROJECT_ROOT))

from src.utils import DATA_PROCESSED_DIR, FIGURES_DIR, REPORTS_DIR  # noqa: E402

STUDENT_NAME = "Emre Kaya"
STUDENT_NUMBER = "231041045"
GITHUB_REPO_URL = "https://github.com/emrekayy/blm308-urun-iade-tahmini"
AI_STATEMENT = "Proje yapılırken yapay zeka araçlarından yararlanılmıştır."

PDF_FONT = "ReportFont"
PDF_FONT_PATHS = [
    "/System/Library/Fonts/Supplemental/Arial.ttf",
    "/Library/Fonts/Arial.ttf",
    "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
]


def load_json(path: Path) -> dict:
    with path.open(encoding="utf-8") as file:
        return json.load(file)


def register_pdf_font() -> str:
    for font_path in PDF_FONT_PATHS:
        if Path(font_path).exists():
            pdfmetrics.registerFont(TTFont(PDF_FONT, font_path))
            return PDF_FONT
    return "Helvetica"


def translate_skipped_chart(text: str) -> str:
    replacements = {
        "Return rate by category (no category column found)": "Kategoriye göre iade oranı (kategori sütunu bulunamadı)",
        "Discount vs return (no discount column found)": "İndirim ve iade ilişkisi (indirim sütunu bulunamadı)",
        "Delivery time vs return (no delivery time column found)": "Teslimat süresi ve iade ilişkisi (teslimat süresi sütunu bulunamadı)",
    }
    return replacements.get(text, text)


def add_heading(document: Document, text: str, level: int = 1) -> None:
    document.add_heading(text, level=level)


def add_paragraph(document: Document, text: str, bold: bool = False) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(11)


def add_bullet_list(document: Document, items: list[str]) -> None:
    for item in items:
        document.add_paragraph(item, style="List Bullet")


def add_comparison_table(document: Document, records: list[dict]) -> None:
    if not records:
        return
    headers = ["Model", "Doğruluk", "Kesinlik", "Duyarlılık", "F1-Skoru", "ROC-AUC"]
    keys = ["Model", "Accuracy", "Precision", "Recall", "F1-Score", "ROC-AUC"]
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = header
    for record in records:
        row_cells = table.add_row().cells
        for index, key in enumerate(keys):
            value = record[key]
            row_cells[index].text = f"{value:.4f}" if isinstance(value, float) else str(value)


def add_figure(document: Document, image_path: Path, caption: str) -> None:
    if image_path.exists():
        document.add_picture(str(image_path), width=Inches(5.5))
        caption_paragraph = document.add_paragraph(caption)
        caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def build_report() -> Path:
    preprocessing = load_json(DATA_PROCESSED_DIR / "preprocessing_metadata.json")
    evaluation = load_json(DATA_PROCESSED_DIR / "evaluation_results.json")

    inspection = preprocessing["inspection"]
    class_dist = inspection["class_distribution"]
    skipped = [translate_skipped_chart(item) for item in preprocessing["eda"]["skipped_charts"]]
    comparison = evaluation["comparison_table"]
    best_model = evaluation["best_model"]
    test_metrics = evaluation["test_metrics"]
    cm = evaluation["confusion_matrix"]

    document = Document()
    title = document.add_heading(
        "BLM308 Veri Madenciliği Final Projesi\n"
        "Çevrimiçi Alışveriş Ürün İade Tahmini",
        level=0,
    )
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_heading(document, "Öğrenci Bilgileri", level=1)
    add_paragraph(document, f"Ad Soyad: {STUDENT_NAME}")
    add_paragraph(document, f"Öğrenci Numarası: {STUDENT_NUMBER}")
    add_paragraph(document, "Ders: BLM308 Veri Madenciliği")
    add_paragraph(document, "Teslim Tarihi: Mayıs 2026")
    add_paragraph(document, f"GitHub Deposu: {GITHUB_REPO_URL}")

    add_heading(document, "Proje Görev Dağılımı", level=1)
    add_paragraph(
        document,
        f"Bu proje bireysel olarak {STUDENT_NAME} tarafından gerçekleştirilmiştir. "
        "Veri ön işleme, keşifsel veri analizi, model eğitimi, değerlendirme, rapor ve "
        "sunum hazırlığının tamamı proje sahibi tarafından yapılmıştır.",
    )

    add_heading(document, "Özet", level=1)
    add_paragraph(
        document,
        "Bu proje, bir e-ticaret siparişinin iade edilip edilmeyeceğini tahmin eden ikili "
        "sınıflandırma problemini ele almaktadır. product_return_prediction.csv veri seti "
        f"({inspection['shape'][0]} kayıt) kullanılarak CRISP-DM metodolojisi uygulanmış, "
        "keşifsel veri analizi gerçekleştirilmiş, fiyat, puan ve ürün kimlik bilgilerinden "
        "özellikler türetilmiş ve beş sınıflandırma algoritması eğitim kümesi üzerinde "
        f"10 katmanlı stratified çapraz doğrulama ile karşılaştırılmıştır. En iyi performansı "
        f"gösteren model {best_model} olmuş; ayrılmış test kümesinde yalnızca bir kez "
        f"değerlendirilmiş ve F1-skoru {test_metrics['f1_score']:.4f}, ROC-AUC değeri "
        f"{test_metrics['roc_auc']:.4f} elde edilmiştir.",
    )

    add_heading(document, "1. Problem Tanımı", level=1)
    add_paragraph(
        document,
        "Ürün iadeleri çevrimiçi perakendede operasyonel maliyet, stok yönetimi sorunları "
        "ve müşteri memnuniyetsizliği yaratmaktadır. Bu projenin amacı, return_status "
        "(0 = iade edilmedi, 1 = iade edildi) değişkenini satın alma öncesinde veya "
        "sonrasında tahmin ederek perakendecilerin kalite kontrolünü önceliklendirmesine, "
        "fiyatlandırma politikalarını düzenlemesine ve müşteri deneyimini iyileştirmesine "
        "yardımcı olmaktır.",
    )

    add_heading(document, "2. Motivasyon ve İş Değeri", level=1)
    add_bullet_list(
        document,
        [
            "Yüksek riskli siparişleri erken tespit ederek ters lojistik maliyetlerini azaltmak.",
            "İade edilecek ürünleri öngörerek envanter planlamasını iyileştirmek.",
            "Düşük puanlı veya yüksek fiyatlı alışverişlerde hedefli müşteri hizmeti sunmak.",
            "Veriye dayalı fiyatlandırma ve ürün portföyü kararlarını desteklemek.",
        ],
    )

    add_heading(document, "3. İlgili Çalışmalar / Literatür", level=1)
    add_paragraph(
        document,
        "E-ticaret analitiğinde iade tahmini; müşteri davranışı, ürün özellikleri ve "
        "satın alma sonrası geri bildirimler kullanılarak incelenmiştir. Önceki çalışmalar, "
        "fiyat duyarlılığı, ürün kalitesi sinyalleri (puanlar) ve ürüne özgü kalıpların "
        "iadeler için güçlü öngörücüler olduğunu göstermektedir. Dengesiz perakende "
        "sınıflandırma görevlerinde ağaç tabanlı topluluk modelleri ve lojistik regresyon "
        "yaygın olarak kullanılan temel yöntemlerdir.",
    )

    add_heading(document, "4. CRISP-DM Metodolojisi", level=1)
    add_bullet_list(
        document,
        [
            "İş Anlayışı: İade tahminini maliyet azaltma odaklı bir kullanım senaryosu olarak tanımlama.",
            "Veri Anlayışı: Şema, eksik değerler ve sınıf dengesizliğini inceleme.",
            "Veri Hazırlığı: Kategorik değişkenleri kodlama, sayısal özellikleri ölçekleme, veriyi bölme.",
            "Modelleme: Decision Tree, Random Forest, Logistic Regression, Naive Bayes, KNN eğitimi.",
            "Değerlendirme: Eğitim kümesinde 10 katmanlı stratified CV; test kümesi yalnızca final değerlendirme.",
            "Dağıtım: Perakende operasyonları için uygulanabilir içgörülerin dokümante edilmesi.",
        ],
    )

    add_heading(document, "5. Veri Seti Kaynağı ve Açıklaması", level=1)
    add_paragraph(document, f"Veri seti dosyası: {preprocessing['dataset_file']}")
    add_paragraph(
        document,
        f"Veri seti {inspection['shape'][0]} gözlem ve {inspection['shape'][1]} sütundan "
        f"oluşmaktadır: {', '.join(inspection['columns'])}.",
    )
    add_paragraph(
        document,
        f"Hedef değişken: {inspection['target_column']} "
        f"(Sınıf 0: {class_dist.get(0, class_dist.get('0', 'N/A'))}, "
        f"Sınıf 1: {class_dist.get(1, class_dist.get('1', 'N/A'))}).",
    )
    add_paragraph(
        document,
        "Eksik değer tespit edilmemiştir. Kategori, indirim ve teslimat süresi sütunları "
        "bu veri setinde bulunmamaktadır.",
    )

    add_heading(document, "6. Keşifsel Veri Analizi (EDA)", level=1)
    add_paragraph(
        document,
        "Sınıf dağılımı, fiyat-iade ilişkisi, puan dağılımı, puan-iade ilişkisi ve "
        "sayısal özellikler arası korelasyon grafikleri oluşturulmuştur.",
    )
    if skipped:
        add_paragraph(document, "Oluşturulmayan grafikler:")
        add_bullet_list(document, skipped)
    add_figure(document, FIGURES_DIR / "01_class_distribution.png", "Şekil 1: Sınıf dağılımı")
    add_figure(document, FIGURES_DIR / "03_price_vs_return.png", "Şekil 2: Fiyat ve iade durumu")
    add_figure(document, FIGURES_DIR / "05_rating_distribution.png", "Şekil 3: Müşteri puanı dağılımı")
    add_figure(document, FIGURES_DIR / "07_correlation_heatmap.png", "Şekil 4: Korelasyon ısı haritası")

    add_heading(document, "7. Ön İşleme Adımları", level=1)
    add_bullet_list(
        document,
        [
            "Tahmin gücü olmayan benzersiz tanımlayıcı order_id sütunu çıkarıldı.",
            "product_id ordinal sayısal özellik olarak kullanılmadı; Label Encoding kaldırıldı.",
            "product_id_frequency: her ürünün yalnızca eğitim kümesindeki görülme sayısı.",
            "product_return_rate: ürün bazlı iade oranı yalnızca eğitim kümesinden hesaplandı (target leakage önlendi).",
            "Eğitimde görülmeyen ürünler için frekans=0, iade oranı=eğitim kümesi genel ortalaması kullanıldı.",
            "price, rating ve türetilen ürün özelliklerine StandardScaler uygulandı (scaler yalnızca eğitim kümesinde fit edildi).",
            "Stratified %80/%20 eğitim-test ayrımı encoding öncesinde yapıldı (random_state=42).",
            "Temizlenmiş veri data/processed/ klasörüne kaydedildi.",
        ],
    )

    add_heading(document, "8. Seçilen Modeller", level=1)
    add_bullet_list(
        document,
        [
            "Decision Tree (Karar Ağacı)",
            "Random Forest (Rastgele Orman)",
            "Logistic Regression (Lojistik Regresyon)",
            "Gaussian Naive Bayes (Naive Bayes)",
            "K-Nearest Neighbors (KNN)",
        ],
    )

    add_heading(document, "9. Değerlendirme Kurulumu", level=1)
    add_paragraph(
        document,
        "Veri seti stratified %80/%20 oranında eğitim (1.200 kayıt) ve test (300 kayıt) "
        "kümelerine ayrılmıştır. product_id için frekans ve iade oranı özellikleri yalnızca "
        "eğitim kümesinden türetilmiştir. Çapraz doğrulama sırasında bu özellikler her fold "
        "içinde yeniden hesaplanarak target leakage önlenmiştir. Model karşılaştırması ve "
        "en iyi model seçimi yalnızca eğitim kümesi üzerinde 10 katmanlı Stratified Cross "
        "Validation ile yapılmıştır. Test kümesi modele hiçbir aşamada dahil edilmemiş; "
        "yalnızca seçilen en iyi model tam eğitim kümesi ile eğitildikten sonra bir kez "
        "final değerlendirme için kullanılmıştır.",
    )
    add_paragraph(
        document,
        "Kullanılan metrikler: Doğruluk (Accuracy), Kesinlik (Precision), Duyarlılık (Recall), "
        "F1-skoru ve ROC-AUC. Model seçim kriteri: çapraz doğrulama ortalama F1-skoru.",
    )

    add_heading(document, "10. Performans Karşılaştırması", level=1)
    add_paragraph(
        document,
        "Aşağıdaki tablo, yalnızca eğitim kümesi üzerinde gerçekleştirilen 10 katmanlı "
        "stratified çapraz doğrulama ortalama sonuçlarını göstermektedir.",
    )
    add_comparison_table(document, comparison)

    add_heading(document, "11. En İyi Model Sonuçları", level=1)
    add_paragraph(document, f"Eğitim kümesi CV F1-skoruna göre seçilen en iyi model: {best_model}")
    add_paragraph(
        document,
        "Seçilen model, tam eğitim kümesi (1.200 kayıt) ile yeniden eğitilmiş ve daha "
        "önce hiç kullanılmamış test kümesi (300 kayıt) üzerinde yalnızca bir kez "
        "değerlendirilmiştir. Aşağıdaki metrikler bu bağımsız final test sonuçlarıdır.",
    )
    add_paragraph(
        document,
        f"Test kümesi metrikleri — Doğruluk: {test_metrics['accuracy']:.4f}, "
        f"Kesinlik: {test_metrics['precision']:.4f}, Duyarlılık: {test_metrics['recall']:.4f}, "
        f"F1-skoru: {test_metrics['f1_score']:.4f}, ROC-AUC: {test_metrics['roc_auc']:.4f}.",
    )
    add_paragraph(
        document,
        f"Confusion matrix (test kümesi): TN={cm[0][0]}, FP={cm[0][1]}, "
        f"FN={cm[1][0]}, TP={cm[1][1]}.",
    )
    add_figure(
        document,
        FIGURES_DIR / "08_confusion_matrix_best_model.png",
        "Şekil 5: En iyi model için karmaşıklık matrisi",
    )
    add_figure(
        document,
        FIGURES_DIR / "09_roc_curve_best_model.png",
        "Şekil 6: En iyi model için ROC eğrisi",
    )
    add_figure(
        document,
        FIGURES_DIR / "10_feature_importance_random_forest.png",
        "Şekil 7: Random Forest özellik önem dereceleri",
    )

    add_heading(document, "12. Önyargı-Varyans Tartışması", level=1)
    add_paragraph(
        document,
        "Karar Ağacı modelleri daha yüksek varyansa sahip olabilir ve ürün düzeyindeki "
        "gürültülü kalıplara aşırı uyum gösterebilir. Random Forest, bagging yöntemiyle "
        "varyansı azaltır ancak yüksek kardinaliteli ürün kimliklerini ezberleyebilir. "
        "Lojistik Regresyon daha düşük varyans ve daha iyi yorumlanabilirlik sunar; fakat "
        "doğrusal olmayan etkileşimleri yeterince yakalayamayabilir. Seçilen model, "
        "görülmemiş siparişler üzerinde genelleme ile tahmin performansı arasında denge kurar.",
    )

    add_heading(document, "13. İş Yorumu", level=1)
    add_paragraph(
        document,
        "Düşük müşteri puanları ve belirli ürün kimlikleri önemli öngörücüler olarak "
        "öne çıkmaktadır. Perakende ekipleri, iade risk skorlarını proaktif destek "
        "süreçlerini tetiklemek, sürekli yüksek iade olasılığına sahip ürünleri incelemek "
        "ve iade oranı yüksek ürünlerin fiyatlandırmasını gözden geçirmek için "
        "kullanabilir.",
    )

    add_heading(document, "14. Kısıtlamalar", level=1)
    add_bullet_list(
        document,
        [
            "Veri setinde müşteri demografisi, kategori, indirim ve teslimat özellikleri yoktur.",
            "Sınıf dengesizliği (~%27 iade) azınlık sınıfının duyarlılığını olumsuz etkileyebilir.",
            "product_id kodlaması yeni ürünler için genellenebilir olmayabilir.",
            "Sonuçlar geçmiş verilere dayanır; mevsimsel etkiler yansıtılmayabilir.",
        ],
    )

    add_heading(document, "15. Gelecek Çalışmalar", level=1)
    add_bullet_list(
        document,
        [
            "Yorum metinleri, teslimat süresi ve promosyon verilerinin entegrasyonu.",
            "Sınıf dengesizliği için SMOTE veya maliyet-duyarlı öğrenme uygulanması.",
            "Label encoding yerine ürün embedding yöntemlerinin kullanılması.",
            "Sipariş akışında gerçek zamanlı skorlama API'si geliştirilmesi.",
        ],
    )

    add_heading(document, "16. Yapay Zeka Kullanım Beyanı", level=1)
    add_paragraph(document, AI_STATEMENT)

    add_heading(document, "17. Kaynaklar", level=1)
    add_bullet_list(
        document,
        [
            "Chapman, P. et al. (2000). CRISP-DM 1.0 Step-by-step data mining guide.",
            "Breiman, L. (2001). Random Forests. Machine Learning, 45(1), 5-32.",
            "Pedregosa, F. et al. (2011). Scikit-learn: Machine Learning in Python.",
        ],
    )

    REPORTS_DIR.mkdir(parents=True, exist_ok=True)
    docx_path = REPORTS_DIR / "BLM308_Final_Report.docx"
    document.save(docx_path)
    return docx_path


def build_pdf_report() -> Path:
    preprocessing = load_json(DATA_PROCESSED_DIR / "preprocessing_metadata.json")
    evaluation = load_json(DATA_PROCESSED_DIR / "evaluation_results.json")

    inspection = preprocessing["inspection"]
    class_dist = inspection["class_distribution"]
    skipped = [translate_skipped_chart(item) for item in preprocessing["eda"]["skipped_charts"]]
    comparison = evaluation["comparison_table"]
    best_model = evaluation["best_model"]
    test_metrics = evaluation["test_metrics"]
    cm = evaluation["confusion_matrix"]

    pdf_path = REPORTS_DIR / "BLM308_Final_Report.pdf"
    REPORTS_DIR.mkdir(parents=True, exist_ok=True)

    font_name = register_pdf_font()
    doc = SimpleDocTemplate(str(pdf_path), pagesize=A4, topMargin=0.7 * inch, bottomMargin=0.7 * inch)
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "TitleCustom", parent=styles["Title"], fontName=font_name, fontSize=16, spaceAfter=16,
    )
    heading_style = ParagraphStyle(
        "HeadingCustom", parent=styles["Heading2"], fontName=font_name, fontSize=13, spaceAfter=8,
    )
    body_style = ParagraphStyle(
        "BodyCustom", parent=styles["BodyText"], fontName=font_name, fontSize=10, leading=14, spaceAfter=6,
    )

    story = []
    story.append(Paragraph("BLM308 Veri Madenciliği Final Projesi", title_style))
    story.append(Paragraph("Çevrimiçi Alışveriş Ürün İade Tahmini", title_style))
    story.append(Spacer(1, 0.2 * inch))

    sections = [
        ("Öğrenci Bilgileri", [
            f"Ad Soyad: {STUDENT_NAME}",
            f"Öğrenci Numarası: {STUDENT_NUMBER}",
            "Ders: BLM308 Veri Madenciliği",
            "Teslim Tarihi: Mayıs 2026",
            f"GitHub Deposu: {GITHUB_REPO_URL}",
        ]),
        ("Proje Görev Dağılımı", [
            f"Bu proje bireysel olarak {STUDENT_NAME} tarafından gerçekleştirilmiştir. "
            "Tüm proje adımları proje sahibi tarafından yapılmıştır.",
        ]),
        ("Özet", [
            f"Bu proje, {inspection['shape'][0]} kayıtlı {preprocessing['dataset_file']} veri seti "
            f"ile e-ticaret ürün iadesi tahmin etmektedir. Beş sınıflandırıcı yalnızca eğitim "
            f"kümesi üzerinde 10 katmanlı stratified çapraz doğrulama ile karşılaştırılmıştır. "
            f"En iyi model: {best_model}. Test kümesi yalnızca bir kez final değerlendirme "
            f"için kullanılmıştır (F1={test_metrics['f1_score']:.4f}, "
            f"ROC-AUC={test_metrics['roc_auc']:.4f}).",
        ]),
        ("1. Problem Tanımı", [
            "return_status (0 = iade yok, 1 = iade var) değişkenini tahmin ederek ters lojistik "
            "maliyetlerini azaltmak ve müşteri deneyimini iyileştirmek.",
        ]),
        ("2. Motivasyon ve İş Değeri", [
            "İade maliyetlerini düşürmek, envanter planlamasını iyileştirmek ve proaktif destek sağlamak.",
        ]),
        ("3. İlgili Çalışmalar / Literatür", [
            "E-ticaret çalışmaları fiyat, puan ve ürün özelliklerinin iade için güçlü öngörücüler olduğunu gösterir.",
        ]),
        ("4. CRISP-DM Metodolojisi", [
            "İş anlayışı, veri anlayışı, hazırlık, modelleme, değerlendirme ve dağıtım aşamaları uygulanmıştır.",
        ]),
        ("5. Veri Seti Kaynağı ve Açıklaması", [
            f"Sütunlar: {', '.join(inspection['columns'])}",
            f"Hedef: {inspection['target_column']} | Sınıf 0: {class_dist.get(0, class_dist.get('0'))}, "
            f"Sınıf 1: {class_dist.get(1, class_dist.get('1'))}",
            "Eksik değer yok. Kategori, indirim ve teslimat sütunları mevcut değil.",
        ]),
        ("6. Keşifsel Veri Analizi", [
            "Sınıf dağılımı, fiyat-iade, puan dağılımı ve korelasyon grafikleri oluşturuldu.",
            f"Oluşturulmayan grafikler: {', '.join(skipped) if skipped else 'Yok'}.",
        ]),
        ("7. Ön İşleme Adımları", [
            "order_id çıkarıldı; product_id ordinal olarak kodlanmadı.",
            "product_id_frequency ve product_return_rate eğitim kümesinden türetildi.",
            "Eğitimde görülmeyen ürünler için frekans=0, iade oranı=eğitim ortalaması.",
            "Scaler yalnızca eğitim kümesinde fit edildi; stratified bölünme encoding öncesinde yapıldı.",
        ]),
        ("8. Seçilen Modeller", [
            "Decision Tree, Random Forest, Logistic Regression, Naive Bayes, KNN.",
        ]),
        ("9. Değerlendirme Kurulumu", [
            "10 katmanlı Stratified CV yalnızca eğitim kümesinde (1.200 kayıt).",
            "Ürün frekans/iade oranı encoding her CV fold içinde yeniden hesaplandı.",
            "Test kümesi (300 kayıt) yalnızca bir kez final değerlendirme için kullanıldı.",
            "Metrikler: Doğruluk, Kesinlik, Duyarlılık, F1, ROC-AUC. Model seçimi: CV F1-skoru.",
        ]),
        ("10. Performans Karşılaştırması", []),
        ("11. En İyi Model Sonuçları", [
            f"En iyi model (eğitim CV): {best_model}",
            "Model tam eğitim kümesi ile eğitildi; test kümesi bağımsız final değerlendirme.",
            f"Test Doğruluk={test_metrics['accuracy']:.4f}, Kesinlik={test_metrics['precision']:.4f}, "
            f"Duyarlılık={test_metrics['recall']:.4f}, F1={test_metrics['f1_score']:.4f}, "
            f"ROC-AUC={test_metrics['roc_auc']:.4f}",
            f"Confusion matrix: TN={cm[0][0]}, FP={cm[0][1]}, FN={cm[1][0]}, TP={cm[1][1]}",
        ]),
        ("12. Önyargı-Varyans Tartışması", [
            "Ağaç modelleri doğrusal olmayan kalıpları yakalar ancak aşırı uyum riski taşır. "
            "Doğrusal modeller daha iyi geneller fakat karmaşık etkileşimleri kaçırabilir.",
        ]),
        ("13. İş Yorumu", [
            "Puan ve ürün bazlı sinyaller proaktif destek ve kalite incelemesi için kullanılabilir.",
        ]),
        ("14. Kısıtlamalar", [
            "Sınırlı özellik seti, sınıf dengesizliği ve product_id genelleme kısıtları.",
        ]),
        ("15. Gelecek Çalışmalar", [
            "Yorum metni, teslimat özellikleri, dengesizlik yönetimi ve gerçek zamanlı dağıtım.",
        ]),
        ("16. Yapay Zeka Kullanım Beyanı", [
            AI_STATEMENT,
        ]),
        ("17. Kaynaklar", [
            "Chapman et al. (2000) CRISP-DM; Breiman (2001) Random Forests; Pedregosa et al. (2011) scikit-learn.",
        ]),
    ]

    for title, bullets in sections:
        story.append(Paragraph(title, heading_style))
        if title == "10. Performans Karşılaştırması":
            table_data = [["Model", "Doğruluk", "Kesinlik", "Duyarlılık", "F1-Skoru", "ROC-AUC"]]
            for row in comparison:
                table_data.append([
                    row["Model"],
                    f"{row['Accuracy']:.4f}",
                    f"{row['Precision']:.4f}",
                    f"{row['Recall']:.4f}",
                    f"{row['F1-Score']:.4f}",
                    f"{row['ROC-AUC']:.4f}",
                ])
            table = Table(table_data, repeatRows=1)
            table.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                ("FONTSIZE", (0, 0), (-1, -1), 8),
                ("FONTNAME", (0, 0), (-1, -1), font_name),
            ]))
            story.append(table)
            story.append(Spacer(1, 0.15 * inch))
            continue

        for bullet in bullets:
            story.append(Paragraph(f"• {bullet}", body_style))
        story.append(Spacer(1, 0.1 * inch))

    for figure_name, caption in [
        ("01_class_distribution.png", "Şekil 1: Sınıf dağılımı"),
        ("03_price_vs_return.png", "Şekil 2: Fiyat ve iade durumu"),
        ("08_confusion_matrix_best_model.png", "Şekil 3: Karmaşıklık matrisi"),
        ("09_roc_curve_best_model.png", "Şekil 4: ROC eğrisi"),
        ("10_feature_importance_random_forest.png", "Şekil 5: Random Forest özellik önem dereceleri"),
    ]:
        image_path = FIGURES_DIR / figure_name
        if image_path.exists():
            story.append(Paragraph(caption, heading_style))
            story.append(Image(str(image_path), width=5.5 * inch, height=3.2 * inch))
            story.append(Spacer(1, 0.1 * inch))

    doc.build(story)
    return pdf_path


def export_pdf(docx_path: Path) -> Path | None:
    return build_pdf_report()


def main() -> None:
    docx_path = build_report()
    print(f"Rapor kaydedildi: {docx_path}")
    pdf_path = export_pdf(docx_path)
    if pdf_path:
        print(f"PDF kaydedildi: {pdf_path}")


if __name__ == "__main__":
    main()
